from __future__ import annotations

import csv
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Iterable
from zipfile import ZIP_DEFLATED, ZipFile

import fitz
import openpyxl
import pikepdf
from PIL import Image
from pypdf import PdfReader, PdfWriter
from reportlab.lib.colors import Color
from reportlab.lib.pagesizes import A4, portrait
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


@dataclass
class PDFBytesResult:
    data: bytes
    filename: str


@dataclass
class FileBytesResult:
    data: bytes
    filename: str
    mime_type: str


def _read(uploaded_file) -> PdfReader:
    return PdfReader(BytesIO(_read_bytes(uploaded_file)))


def _read_bytes(uploaded_file) -> bytes:
    uploaded_file.seek(0)
    return uploaded_file.read()


def parse_page_list(raw: str) -> list[int]:
    pages: list[int] = []
    for chunk in raw.split(","):
        chunk = chunk.strip()
        if not chunk:
            continue
        if "-" in chunk:
            start_s, end_s = chunk.split("-", 1)
            start, end = int(start_s), int(end_s)
            if start <= 0 or end <= 0 or end < start:
                raise ValueError(f"Invalid page range: {chunk}")
            pages.extend(range(start, end + 1))
        else:
            value = int(chunk)
            if value <= 0:
                raise ValueError(f"Invalid page number: {chunk}")
            pages.append(value)
    return pages


def parse_page_ranges(raw: str) -> list[tuple[int, int]]:
    ranges: list[tuple[int, int]] = []
    for chunk in raw.split(","):
        chunk = chunk.strip()
        if not chunk:
            continue
        if "-" in chunk:
            start_s, end_s = chunk.split("-", 1)
            start, end = int(start_s), int(end_s)
        else:
            start = end = int(chunk)
        if start <= 0 or end <= 0 or end < start:
            raise ValueError(f"Invalid page range: {chunk}")
        ranges.append((start, end))
    return ranges


def merge_pdfs(files: list) -> PDFBytesResult:
    writer = PdfWriter()
    for file in files:
        reader = _read(file)
        for page in reader.pages:
            writer.add_page(page)
    buffer = BytesIO()
    writer.write(buffer)
    return PDFBytesResult(buffer.getvalue(), "merged.pdf")


def split_pdf(file, page_ranges: str) -> bytes:
    reader = _read(file)
    zip_buffer = BytesIO()
    with ZipFile(zip_buffer, "w", ZIP_DEFLATED) as zf:
        for idx, (start, end) in enumerate(parse_page_ranges(page_ranges), start=1):
            writer = PdfWriter()
            for page_num in range(start - 1, min(end, len(reader.pages))):
                writer.add_page(reader.pages[page_num])
            buf = BytesIO()
            writer.write(buf)
            zf.writestr(f"split_{idx}_{start}-{end}.pdf", buf.getvalue())
    return zip_buffer.getvalue()


def delete_pages(file, pages_to_delete: Iterable[int]) -> PDFBytesResult:
    reader = _read(file)
    writer = PdfWriter()
    drop = {p - 1 for p in pages_to_delete}
    for i, page in enumerate(reader.pages):
        if i not in drop:
            writer.add_page(page)
    buf = BytesIO()
    writer.write(buf)
    return PDFBytesResult(buf.getvalue(), "pages_deleted.pdf")


def rotate_pdf(file, pages: Iterable[int], angle: int) -> PDFBytesResult:
    reader = _read(file)
    writer = PdfWriter()
    targets = {p - 1 for p in pages}
    for i, page in enumerate(reader.pages):
        if i in targets:
            page.rotate(angle)
        writer.add_page(page)
    buf = BytesIO()
    writer.write(buf)
    return PDFBytesResult(buf.getvalue(), "rotated.pdf")


def reorder_pdf(file, order: list[int]) -> PDFBytesResult:
    reader = _read(file)
    total = len(reader.pages)
    if sorted(order) != list(range(1, total + 1)):
        raise ValueError("Page order must include every page exactly once.")
    writer = PdfWriter()
    for page_number in order:
        writer.add_page(reader.pages[page_number - 1])
    buf = BytesIO()
    writer.write(buf)
    return PDFBytesResult(buf.getvalue(), "reordered.pdf")


def compress_pdf(file) -> PDFBytesResult:
    source = _read_bytes(file)
    input_buffer = BytesIO(source)
    output_buffer = BytesIO()

    with pikepdf.open(input_buffer) as pdf:
        pdf.remove_unreferenced_resources()
        pdf.save(
            output_buffer,
            compress_streams=True,
            recompress_flate=True,
            object_stream_mode=pikepdf.ObjectStreamMode.generate,
            linearize=True,
        )

    result = output_buffer.getvalue()
    if len(result) >= len(source):
        result = source
    return PDFBytesResult(result, "compressed.pdf")


def pdf_to_images(file, image_format: str = "png", dpi: int = 150) -> bytes:
    source = _read_bytes(file)
    zoom = max(dpi, 72) / 72
    matrix = fitz.Matrix(zoom, zoom)
    image_format = image_format.lower()
    ext = "jpg" if image_format in {"jpg", "jpeg"} else "png"

    zip_buffer = BytesIO()
    with fitz.open(stream=source, filetype="pdf") as doc, ZipFile(zip_buffer, "w", ZIP_DEFLATED) as zf:
        for index, page in enumerate(doc, start=1):
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            img_bytes = pix.tobytes("jpeg" if ext == "jpg" else "png")
            zf.writestr(f"page_{index:03d}.{ext}", img_bytes)
    return zip_buffer.getvalue()


def images_to_pdf(files: list, output_name: str = "images_to_pdf.pdf") -> PDFBytesResult:
    if not files:
        raise ValueError("Upload at least one image.")

    images: list[Image.Image] = []
    for file in files:
        image = Image.open(BytesIO(file.getvalue()))
        if image.mode in ("RGBA", "LA"):
            background = Image.new("RGB", image.size, "white")
            background.paste(image, mask=image.getchannel("A"))
            image = background
        else:
            image = image.convert("RGB")
        images.append(image)

    first, *rest = images
    buffer = BytesIO()
    first.save(buffer, format="PDF", save_all=True, append_images=rest)
    return PDFBytesResult(buffer.getvalue(), output_name)


def _register_arabic_font(font_path: str | None = None) -> str:
    candidates = [
        font_path,
        "/usr/share/fonts/truetype/noto/NotoNaskhArabic-Regular.ttf",
        "/usr/share/fonts/opentype/fonts-hosny-amiri/Amiri-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansArabic-Regular.ttf",
    ]
    for path in candidates:
        if path and Path(path).exists():
            try:
                pdfmetrics.registerFont(TTFont("ArabicWatermarkFont", path))
            except Exception:
                pass
            return "ArabicWatermarkFont"
    return "Helvetica"


def add_text_watermark(
    file,
    text: str,
    pages: Iterable[int] | None = None,
    opacity: float = 0.18,
    rotation: int = 35,
    font_size: int = 42,
    font_path: str | None = None,
) -> PDFBytesResult:
    source = _read_bytes(file)
    reader = PdfReader(BytesIO(source))
    writer = PdfWriter()
    selected = {p - 1 for p in pages} if pages else set(range(len(reader.pages)))
    font_name = _register_arabic_font(font_path)

    for index, page in enumerate(reader.pages):
        if index not in selected:
            writer.add_page(page)
            continue

        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        overlay_buffer = BytesIO()
        c = canvas.Canvas(overlay_buffer, pagesize=portrait((width, height)))
        c.saveState()
        c.setFillColor(Color(0.12, 0.16, 0.28, alpha=max(0.02, min(opacity, 0.95))))
        c.setFont(font_name, font_size)
        c.translate(width / 2, height / 2)
        c.rotate(rotation)
        c.drawCentredString(0, 0, text)
        c.restoreState()
        c.save()
        overlay_buffer.seek(0)
        overlay_pdf = PdfReader(overlay_buffer)
        page.merge_page(overlay_pdf.pages[0])
        writer.add_page(page)

    out = BytesIO()
    writer.write(out)
    return PDFBytesResult(out.getvalue(), "watermarked.pdf")


def page_count(file) -> int:
    return len(_read(file).pages)


def _text_to_pdf_bytes(text: str, title: str = "Document") -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    text_obj = c.beginText(40, height - 60)
    text_obj.setFont("Helvetica", 11)
    line_height = 16
    for raw_line in (text or "").splitlines() or [""]:
        line = raw_line.strip() or " "
        while len(line) > 95:
            text_obj.textLine(line[:95])
            line = line[95:]
            if text_obj.getY() < 60:
                c.drawText(text_obj)
                c.showPage()
                text_obj = c.beginText(40, height - 60)
                text_obj.setFont("Helvetica", 11)
        text_obj.textLine(line)
        if text_obj.getY() < 60:
            c.drawText(text_obj)
            c.showPage()
            text_obj = c.beginText(40, height - 60)
            text_obj.setFont("Helvetica", 11)
    c.drawText(text_obj)
    c.setTitle(title)
    c.save()
    return buffer.getvalue()


def word_to_pdf(file) -> PDFBytesResult:
    suffix = Path(file.name).suffix.lower()
    source = _read_bytes(file)
    if suffix == ".docx":
        try:
            from docx import Document

            with tempfile.TemporaryDirectory(prefix="word-to-pdf-") as tmp:
                docx_path = Path(tmp) / "input.docx"
                docx_path.write_bytes(source)
                doc = Document(str(docx_path))
                text = "\n".join(p.text for p in doc.paragraphs)
                return PDFBytesResult(_text_to_pdf_bytes(text, title=file.name), f"{Path(file.name).stem}.pdf")
        except Exception as exc:
            raise RuntimeError(f"Word to PDF failed: {exc}") from exc

    if suffix == ".txt":
        text = source.decode("utf-8", errors="ignore")
        return PDFBytesResult(_text_to_pdf_bytes(text, title=file.name), f"{Path(file.name).stem}.pdf")

    raise RuntimeError("Upload a .docx or .txt file for Word to PDF.")


def pdf_to_word(file) -> FileBytesResult:
    source = _read_bytes(file)
    text_pages: list[str] = []
    with fitz.open(stream=source, filetype="pdf") as doc:
        for i, page in enumerate(doc, start=1):
            text_pages.append(f"Page {i}\n{'=' * 20}\n{page.get_text('text').strip()}\n")
    text = "\n".join(text_pages).strip()
    try:
        from docx import Document

        docx = Document()
        docx.add_heading(Path(file.name).stem, level=1)
        for chunk in text_pages:
            for line in chunk.splitlines():
                docx.add_paragraph(line)
        buffer = BytesIO()
        docx.save(buffer)
        return FileBytesResult(buffer.getvalue(), f"{Path(file.name).stem}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception:
        return FileBytesResult(text.encode("utf-8"), f"{Path(file.name).stem}.txt", "text/plain")


def pdf_to_excel(file) -> FileBytesResult:
    source = _read_bytes(file)
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "PDF Content"
    sheet.append(["Page", "Line", "Text"])
    with fitz.open(stream=source, filetype="pdf") as doc:
        for page_number, page in enumerate(doc, start=1):
            lines = [line.strip() for line in page.get_text("text").splitlines() if line.strip()]
            if not lines:
                sheet.append([page_number, 1, ""])
            for line_number, line in enumerate(lines, start=1):
                sheet.append([page_number, line_number, line])
    buffer = BytesIO()
    workbook.save(buffer)
    return FileBytesResult(buffer.getvalue(), f"{Path(file.name).stem}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


def excel_to_pdf(file) -> PDFBytesResult:
    workbook = openpyxl.load_workbook(BytesIO(_read_bytes(file)), data_only=True)
    rows: list[str] = []
    for sheet in workbook.worksheets:
        rows.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(v) for v in row if v is not None and str(v).strip()]
            rows.append(" | ".join(values))
        rows.append("")
    return PDFBytesResult(_text_to_pdf_bytes("\n".join(rows), title=file.name), f"{Path(file.name).stem}.pdf")


def tool_availability() -> dict[str, bool]:
    return {
        "libreoffice": bool(shutil.which("libreoffice") or shutil.which("soffice")),
        "tesseract": bool(shutil.which("tesseract")),
        "ghostscript": bool(shutil.which("gs")),
    }
